import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import markdown2
import requests
from PIL import Image as PILImage
import os
import tempfile
import re
from datetime import datetime, timedelta

# Register Unicode font
FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")
if os.path.exists(FONT_PATH):
    pdfmetrics.registerFont(TTFont("DejaVuSans", FONT_PATH))

# Template path
TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), 
    "..", 
    "templates", 
    "Beige and Blue Minimal and Professional Daily Travel Itinerary Planner.docx"
)

def fetch_static_map(places, width=600, height=350):
    marker_strs = []
    for i, place in enumerate(places):
        lat, lon = place.get("coords", [None, None])
        if lat is not None and lon is not None:
            color = ["red", "blue", "green", "yellow", "purple"][i % 5]
            marker_strs.append(f"markers={lat},{lon},{color}{i+1}")
    markers = "&".join(marker_strs)
    center = f"{places[0]['coords'][0]},{places[0]['coords'][1]}"
    url = f"https://staticmap.openstreetmap.de/staticmap.php?center={center}&zoom=13&size={width}x{height}&{markers}"
    try:
        resp = requests.get(url, timeout=5)
        if resp.status_code == 200:
            return io.BytesIO(resp.content)
    except Exception as e:
        print(f"Could not fetch static map: {e}")
    return None

def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """Replace text in a paragraph while preserving formatting"""
    if placeholder in paragraph.text:
        # Get all runs
        runs = paragraph.runs
        # Find which runs contain the placeholder
        for i, run in enumerate(runs):
            if placeholder in run.text:
                # Replace in this run
                run.text = run.text.replace(placeholder, replacement)
            elif paragraph.text.find(placeholder) != -1:
                # Placeholder spans multiple runs, need to reconstruct
                full_text = paragraph.text
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, replacement)
                    # Clear all runs
                    for run in runs:
                        run.text = ""
                    # Add new text to first run
                    if runs:
                        runs[0].text = new_text
                    break

def replace_placeholders_in_docx(doc, placeholders):
    """Replace placeholders in all paragraphs, tables, and headers/footers"""
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            replace_text_in_paragraph(paragraph, placeholder, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        replace_text_in_paragraph(paragraph, placeholder, str(value))
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for placeholder, value in placeholders.items():
                    replace_text_in_paragraph(paragraph, placeholder, str(value))
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for placeholder, value in placeholders.items():
                    replace_text_in_paragraph(paragraph, placeholder, str(value))

def add_itinerary_to_template(doc, markdown_text):
    """Add itinerary content to the template document"""
    # Parse markdown and add structured content
    html = markdown2.markdown(markdown_text)
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        
        # Find a good place to insert content or add at the end
        # Add a page break before itinerary content
        doc.add_page_break()
        
        # Add detailed itinerary heading
        doc.add_heading('Detailed Itinerary', level=1)
        
        for elem in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
            if elem.name == "h1":
                doc.add_heading(elem.get_text(), level=1)
            elif elem.name == "h2":
                doc.add_heading(elem.get_text(), level=2)
            elif elem.name == "h3":
                doc.add_heading(elem.get_text(), level=3)
            elif elem.name == "ul":
                for li in elem.find_all("li"):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            elif elem.name == "ol":
                for li in elem.find_all("li"):
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif elem.name == "p" and elem.get_text().strip():
                doc.add_paragraph(elem.get_text())
    except ImportError:
        # Fallback if BeautifulSoup is not available
        doc.add_page_break()
        doc.add_heading('Detailed Itinerary', level=1)
        # Split by lines and add as paragraphs
        lines = markdown_text.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('#'):
                    # It's a heading
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('# ').strip()
                    doc.add_heading(title, level=min(level, 3))
                else:
                    doc.add_paragraph(line)

def create_itinerary_docx(markdown_text, places=None, options=None):
    """Create DOCX using the template"""
    
    # Load the template
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
        print(f"Using template: {TEMPLATE_PATH}")
    else:
        print(f"Template not found at {TEMPLATE_PATH}, creating new document")
        doc = Document()
        # Add basic structure if no template
        doc.add_heading('Travel Itinerary', 0)
    
    # Extract destination from places
    destination = "Your Destination"
    if places and len(places) > 0:
        # Get the most common city from places
        cities = []
        for place in places:
            name = place.get('name', '')
            if ',' in name:
                city = name.split(',')[-1].strip()
                cities.append(city)
        
        if cities:
            # Use the most frequent city
            from collections import Counter
            destination = Counter(cities).most_common(1)[0][0]
        else:
            destination = places[0].get('name', 'Your Destination')
    
    # Create date range
    start_date = datetime.now()
    days = 3  # default
    if options and 'days' in options:
        try:
            days = int(options['days'])
        except (ValueError, TypeError):
            days = 3
    
    end_date = start_date + timedelta(days=days-1)
    date_range = f"{start_date.strftime('%B %d')} - {end_date.strftime('%B %d, %Y')}"
    
    # Budget and people info
    budget = "Not specified"
    people = "1"
    if options:
        if options.get('budget'):
            budget = f"₹{options['budget']}"
        if options.get('people'):
            people = str(options['people'])
    
    # Prepare placeholders for replacement
    placeholders = {
        '[DESTINATION]': destination,
        '[DATE_RANGE]': date_range,
        '[TRAVELER_NAME]': 'Traveler',
        '[CURRENT_DATE]': datetime.now().strftime('%B %d, %Y'),
        '[BUDGET]': budget,
        '[PEOPLE]': people,
        '[DAYS]': str(days),
        # Add more common placeholders
        'DESTINATION': destination,
        'DATE_RANGE': date_range,
        'TRAVELER_NAME': 'Traveler',
        'CURRENT_DATE': datetime.now().strftime('%B %d, %Y'),
        'BUDGET': budget,
        'PEOPLE': people,
        'DAYS': str(days),
    }
    
    print(f"Replacing placeholders: {placeholders}")
    
    # Replace placeholders
    replace_placeholders_in_docx(doc, placeholders)
    
    # Add detailed itinerary content
    add_itinerary_to_template(doc, markdown_text)
    
    # Try to add map image
    if places:
        map_img_io = fetch_static_map(places)
        if map_img_io:
            try:
                # Save image to temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    pil_img = PILImage.open(map_img_io)
                    pil_img.thumbnail((400, 250))
                    pil_img.save(tmp_file.name, format='PNG')
                    
                    # Add map section
                    doc.add_heading('Route Map', level=2)
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(tmp_file.name, width=Inches(5))
                    
                # Clean up temporary file
                os.unlink(tmp_file.name)
            except Exception as e:
                print(f"Could not add map image to DOCX: {e}")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_itinerary_pdf(markdown_text, places=None):
    """Keep existing PDF function for backward compatibility"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    
    # Use DejaVuSans if available
    if os.path.exists(FONT_PATH):
        for style_name in styles.byName:
            styles[style_name].fontName = "DejaVuSans"
        styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Heading1'], alignment=TA_CENTER, fontName="DejaVuSans"))
    else:
        styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Heading1'], alignment=TA_CENTER))
    
    story = []

    story.append(Paragraph("🗺️ Your Bagpack Itinerary", styles['CenterTitle']))
    story.append(Spacer(1, 18))

    # Try to add map image
    if places:
        map_img_io = fetch_static_map(places)
        if map_img_io:
            try:
                pil_img = PILImage.open(map_img_io)
                pil_img.thumbnail((400, 250))
                img_io = io.BytesIO()
                pil_img.save(img_io, format='PNG')
                img_io.seek(0)
                story.append(Image(img_io, width=pil_img.width, height=pil_img.height))
                story.append(Spacer(1, 12))
            except Exception as e:
                print(f"Could not add map image to PDF: {e}")

    html = markdown2.markdown(markdown_text)
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for elem in soup.contents:
            if elem and hasattr(elem, 'name'):
                if elem.name == "h1":
                    story.append(Paragraph(f"<b>{elem.get_text()}</b>", styles['Heading1']))
                elif elem.name == "h2":
                    story.append(Paragraph(f"<b>{elem.get_text()}</b>", styles['Heading2']))
                elif elem.name == "h3":
                    story.append(Paragraph(f"<b>{elem.get_text()}</b>", styles['Heading3']))
                elif elem.name == "ul":
                    for li in elem.find_all("li"):
                        story.append(Paragraph(f"• {li.get_text()}", styles['Normal']))
                elif elem.name == "ol":
                    for idx, li in enumerate(elem.find_all("li"), 1):
                        story.append(Paragraph(f"{idx}. {li.get_text()}", styles['Normal']))
                elif elem.name == "p":
                    story.append(Paragraph(elem.get_text(), styles['Normal']))
                story.append(Spacer(1, 6))
    except ImportError:
        # Fallback without BeautifulSoup
        lines = markdown_text.split('\n')
        for line in lines:
            story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer